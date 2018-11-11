import json
import os
from pyquery import PyQuery as pq
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_chapters(url):
    '''获取章节列表，包括各章节的标题，url，并预留正文空白'''
    doc = pq(url, encoding='gb2312')
    # 最后一项id=d999，非小说相关内容需排除，使用css的:not()方法排除
    chapters = [{
        'title': item.text(),
        'url': url + item.attr('href'),
        'content': ''
    } for item in doc('table.acss td.ccss a:not(#d999)').items()]
    return chapters


def insert_cotent_to_chapters(chapters):
    '''将正文内容插入章节列表中'''
    for chapter in chapters:
        content = get_content(chapter['url'])
        chapter['content'] = content
        print('%s 读取完毕！' % chapter['title'])


def get_content(url):
    '''获取正文内容'''
    # print(url)
    doc = pq(url, encoding='gb2312')
    content_item = doc('#content')
    content = content_item.text()
    return content


def save_file_txt(chapters):
    '''保存为txt文件'''
    with open('九阴九阳.txt', 'w', encoding='utf-8') as file:
        for chapter in chapters:
            file.write(chapter['title'] + '\n')
            file.write(chapter['content'] + '\n')
            file.write('\n\n\n===================================\n')
            print('%s 保存完毕！' % chapter['title'])


def save_file_docx(filename_json):
    '''保存为docx文件'''
    document = Document()

    # 设置普通段落的默认格式
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.styles['Normal'].paragraph_format.first_line_indent = Pt(48)

    # 添加标题，并设置为居中对齐
    p = document.add_heading('九阴九阳', 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.paragraph_format.alignment = 1 # 这样也可以
    # 0-左对齐 1-居中 2-右对齐 3-两端对齐 4-分散对齐 5678不知道是什么

    with open(filename_json, encoding='utf-8') as file:
        chapters = json.load(file)

    # 读取每一章节，添加至docx，每章节之后插入分页符
    for chapter in chapters:
        document.add_heading(chapter['title'], 2)
        for paragraph in chapter['content'].split('\n\n'):
            p = document.add_paragraph(paragraph.replace('�', ''))
            p.paragraph_format.first_line_indent = Pt(24)
        document.add_page_break()
        print('%s 保存完毕！' % chapter['title'])

    document.save('九阴九阳.docx')


def save_file_json(chapters):
    with open('九阴九阳.json', 'w', encoding='utf-8') as file:
        json.dump(chapters, file, ensure_ascii=False, indent=2)
    print('小说全部读取完毕，准备保存')


def main():
    filename_json = '九阴九阳.json'
    if not os.path.exists(filename_json):
        url = 'https://www.555zw.com/book/7/7509/'
        chapters = get_chapters(url)
        insert_cotent_to_chapters(chapters)
        save_file_json(chapters)

    # save_file_txt(chapters)
    save_file_docx(filename_json)
    print('小说全部保存完毕！')


if __name__ == '__main__':
    main()
